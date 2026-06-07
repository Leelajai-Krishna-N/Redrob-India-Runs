from __future__ import annotations

import json
import math
import re
from collections import Counter, defaultdict
from datetime import date
from itertools import combinations
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from scipy.stats import entropy
from sklearn.cluster import AgglomerativeClustering, KMeans
from sklearn.decomposition import PCA
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics import silhouette_score
from sklearn.preprocessing import StandardScaler
from tqdm import tqdm


def find_project_root() -> Path:
    for parent in Path(__file__).resolve().parents:
        if (parent / "data").exists() and (parent / "src").exists():
            return parent
    return Path(__file__).resolve().parents[2]


ROOT = find_project_root()
DATA_DIR = (
    ROOT
    / "data"
    / "raw"
    / "[PUB] India_runs_data_and_ai_challenge"
    / "[PUB] India_runs_data_and_ai_challenge"
    / "India_runs_data_and_ai_challenge"
)
DATASET_PATH = DATA_DIR / "candidates.jsonl"
JD_PATH = DATA_DIR / "job_description.docx"
OUTPUT_DIR = ROOT / "outputs" / "eda"
VIZ_DIR = OUTPUT_DIR / "visualizations"
SAMPLES_DIR = OUTPUT_DIR / "profile_samples"
RANDOM_SEED = 42
TODAY = pd.Timestamp(date.today())


JD_TERMS = {
    "retrieval": [
        "retrieval",
        "information retrieval",
        "semantic search",
        "rag",
        "bm25",
        "hybrid search",
        "query",
    ],
    "ranking": [
        "ranking",
        "ranker",
        "learning to rank",
        "ltr",
        "relevance",
        "rerank",
        "reranker",
    ],
    "recommendation": [
        "recommendation",
        "recommender",
        "personalization",
        "collaborative filtering",
        "matching",
    ],
    "search_infrastructure": [
        "search",
        "elasticsearch",
        "opensearch",
        "solr",
        "indexing",
        "search infrastructure",
    ],
    "embeddings": [
        "embedding",
        "embeddings",
        "sentence transformers",
        "vector",
        "ann",
        "faiss",
    ],
    "vector_databases": [
        "vector database",
        "vector db",
        "pinecone",
        "weaviate",
        "milvus",
        "qdrant",
        "chromadb",
    ],
    "evaluation": [
        "evaluation",
        "eval",
        "ndcg",
        "mrr",
        "precision",
        "recall",
        "offline evaluation",
        "ab testing",
        "a/b testing",
    ],
}


PRODUCT_COMPANIES = {
    "google",
    "microsoft",
    "amazon",
    "meta",
    "facebook",
    "netflix",
    "apple",
    "uber",
    "airbnb",
    "linkedin",
    "adobe",
    "salesforce",
    "atlassian",
    "stripe",
    "spotify",
    "flipkart",
    "swiggy",
    "zomato",
    "razorpay",
    "paytm",
    "phonepe",
    "zoho",
    "freshworks",
}
SERVICE_COMPANIES = {
    "tcs",
    "infosys",
    "wipro",
    "cognizant",
    "accenture",
    "capgemini",
    "mindtree",
    "ltimindtree",
    "hcl",
    "tech mahindra",
    "deloitte",
    "ey",
    "pwc",
    "kpmg",
}


def ensure_dirs() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    VIZ_DIR.mkdir(exist_ok=True)
    SAMPLES_DIR.mkdir(exist_ok=True)


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).lower()
    text = re.sub(r"[^a-z0-9+#./ -]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def tokenize(value: Any) -> list[str]:
    return re.findall(r"[a-z0-9+#.]+", normalize_text(value))


def load_candidates(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in tqdm(handle, desc="Loading candidates", unit="profile"):
            if line.strip():
                rows.append(json.loads(line))
    return rows


def read_docx_text(path: Path) -> str:
    if not path.exists():
        return ""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def value_is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, float) and math.isnan(value):
        return True
    if isinstance(value, str) and value.strip() == "":
        return True
    if isinstance(value, (list, dict)) and len(value) == 0:
        return True
    return False


def flatten_for_audit(obj: Any, prefix: str = "") -> dict[str, Any]:
    out: dict[str, Any] = {}
    if isinstance(obj, dict):
        for key, value in obj.items():
            path = f"{prefix}.{key}" if prefix else key
            if isinstance(value, list):
                out[f"{path}.__count"] = len(value)
                if value and isinstance(value[0], dict):
                    child_keys = sorted({k for item in value if isinstance(item, dict) for k in item})
                    for child_key in child_keys:
                        out[f"{path}.[].{child_key}.__fill_count"] = sum(
                            not value_is_empty(item.get(child_key)) for item in value if isinstance(item, dict)
                        )
                else:
                    out[path] = value
            elif isinstance(value, dict):
                out.update(flatten_for_audit(value, path))
            else:
                out[path] = value
    return out


def compute_entropy(values: list[Any]) -> float:
    counts = Counter(json.dumps(v, sort_keys=True, default=str) for v in values if not value_is_empty(v))
    if not counts:
        return 0.0
    return float(entropy(list(counts.values()), base=2))


def schema_audit(candidates: list[dict[str, Any]]) -> pd.DataFrame:
    sample = pd.Series(candidates).sample(n=min(1000, len(candidates)), random_state=RANDOM_SEED).tolist()
    flattened = [flatten_for_audit(row) for row in sample]
    keys = sorted({key for row in flattened for key in row})
    rows = []
    for key in keys:
        values = [row.get(key) for row in flattened]
        present = [not value_is_empty(v) for v in values]
        non_empty = [v for v, ok in zip(values, present) if ok]
        unique_count = len({json.dumps(v, sort_keys=True, default=str) for v in non_empty})
        most_common = Counter(json.dumps(v, sort_keys=True, default=str) for v in non_empty).most_common(1)
        rows.append(
            {
                "field": key,
                "sample_size": len(sample),
                "fill_rate": round(sum(present) / len(sample), 4),
                "empty_rate": round(1 - (sum(present) / len(sample)), 4),
                "unique_count": unique_count,
                "uniqueness_rate": round(unique_count / max(1, len(non_empty)), 4),
                "entropy_bits": round(compute_entropy(values), 4),
                "most_common_value": most_common[0][0][:250] if most_common else "",
                "most_common_share": round(most_common[0][1] / max(1, len(non_empty)), 4) if most_common else 0.0,
                "audit_note": audit_note(sum(present) / len(sample), unique_count, len(non_empty), most_common),
            }
        )
    df = pd.DataFrame(rows).sort_values(["fill_rate", "field"], ascending=[True, True])
    df.to_csv(OUTPUT_DIR / "schema_audit.csv", index=False)
    return df


def audit_note(fill_rate: float, unique_count: int, non_empty_count: int, most_common: list[tuple[str, int]]) -> str:
    notes = []
    if fill_rate < 0.05:
        notes.append("almost always empty")
    elif fill_rate < 0.5:
        notes.append("often empty")
    if non_empty_count and unique_count == 1:
        notes.append("single repeated value")
    if most_common and non_empty_count and most_common[0][1] / non_empty_count > 0.95 and unique_count > 1:
        notes.append("dominant repeated value")
    if non_empty_count and unique_count / non_empty_count > 0.95 and non_empty_count > 50:
        notes.append("nearly unique per profile")
    return "; ".join(notes)


def build_frames(candidates: list[dict[str, Any]]) -> dict[str, pd.DataFrame]:
    profiles = []
    careers = []
    skills = []
    education = []
    certs = []
    signals = []
    raw_texts = []
    for cand in candidates:
        cid = cand["candidate_id"]
        profile = cand.get("profile", {})
        profiles.append({"candidate_id": cid, **profile})
        text_parts = [profile.get("headline", ""), profile.get("summary", ""), profile.get("current_title", "")]
        for idx, role in enumerate(cand.get("career_history", [])):
            careers.append({"candidate_id": cid, "role_index": idx, **role})
            text_parts.extend([role.get("title", ""), role.get("description", ""), role.get("company", "")])
        for idx, skill in enumerate(cand.get("skills", [])):
            skills.append({"candidate_id": cid, "skill_index": idx, **skill})
            text_parts.append(skill.get("name", ""))
        for idx, edu in enumerate(cand.get("education", [])):
            education.append({"candidate_id": cid, "education_index": idx, **edu})
        for idx, cert in enumerate(cand.get("certifications", [])):
            certs.append({"candidate_id": cid, "certification_index": idx, **cert})
        signal = cand.get("redrob_signals", {}).copy()
        salary = signal.pop("expected_salary_range_inr_lpa", {}) or {}
        assessments = signal.pop("skill_assessment_scores", {}) or {}
        signal["expected_salary_min_lpa"] = salary.get("min")
        signal["expected_salary_max_lpa"] = salary.get("max")
        signal["skill_assessment_count"] = len(assessments)
        signal["skill_assessment_mean"] = float(np.mean(list(assessments.values()))) if assessments else np.nan
        signals.append({"candidate_id": cid, **signal})
        raw_texts.append({"candidate_id": cid, "search_text": normalize_text(" ".join(text_parts))})
    frames = {
        "profiles": pd.DataFrame(profiles),
        "careers": pd.DataFrame(careers),
        "skills": pd.DataFrame(skills),
        "education": pd.DataFrame(education),
        "certs": pd.DataFrame(certs),
        "signals": pd.DataFrame(signals),
        "texts": pd.DataFrame(raw_texts),
    }
    for date_col in ["signup_date", "last_active_date"]:
        if date_col in frames["signals"]:
            frames["signals"][date_col] = pd.to_datetime(frames["signals"][date_col], errors="coerce")
    for date_col in ["start_date", "end_date"]:
        if date_col in frames["careers"]:
            frames["careers"][date_col] = pd.to_datetime(frames["careers"][date_col], errors="coerce")
    return frames


def save_bar(counter: pd.Series, title: str, filename: str, top_n: int = 25) -> None:
    data = counter.head(top_n)
    if data.empty:
        return
    plt.figure(figsize=(12, max(5, min(10, len(data) * 0.35))))
    sns.barplot(x=data.values, y=data.index, color="#2f6f8f")
    plt.title(title)
    plt.xlabel("Count")
    plt.ylabel("")
    plt.tight_layout()
    plt.savefig(VIZ_DIR / filename, dpi=160)
    plt.close()


def save_hist(series: pd.Series, title: str, filename: str, bins: int = 40, kde: bool = False) -> None:
    clean = pd.to_numeric(series, errors="coerce").dropna()
    if clean.empty:
        return
    plt.figure(figsize=(9, 5))
    sns.histplot(clean, bins=bins, kde=kde, color="#3f7f5f")
    plt.title(title)
    plt.tight_layout()
    plt.savefig(VIZ_DIR / filename, dpi=160)
    plt.close()


def save_box(series: pd.Series, title: str, filename: str) -> None:
    clean = pd.to_numeric(series, errors="coerce").dropna()
    if clean.empty:
        return
    plt.figure(figsize=(8, 3))
    sns.boxplot(x=clean, color="#a0643c")
    plt.title(title)
    plt.tight_layout()
    plt.savefig(VIZ_DIR / filename, dpi=160)
    plt.close()


def dataset_overview(frames: dict[str, pd.DataFrame], total: int) -> pd.DataFrame:
    profiles = frames["profiles"]
    nested = pd.DataFrame({"candidate_id": profiles["candidate_id"]})
    nested["career_history_count"] = frames["careers"].groupby("candidate_id").size().reindex(profiles["candidate_id"], fill_value=0).values
    nested["skills_count"] = frames["skills"].groupby("candidate_id").size().reindex(profiles["candidate_id"], fill_value=0).values
    nested["certifications_count"] = frames["certs"].groupby("candidate_id").size().reindex(profiles["candidate_id"], fill_value=0).values
    nested["education_count"] = frames["education"].groupby("candidate_id").size().reindex(profiles["candidate_id"], fill_value=0).values

    rows = [{"section": "dataset", "metric": "candidate_count", "value": total}]
    for col in profiles.columns:
        rows.append({"section": "profile_missing", "metric": col, "value": int(profiles[col].isna().sum())})
    for col in ["years_of_experience", "current_company_size", "current_industry", "country"]:
        counts = profiles[col].value_counts(dropna=False).head(100)
        for key, value in counts.items():
            rows.append({"section": f"distribution_{col}", "metric": str(key), "value": int(value)})
    for col in ["career_history_count", "skills_count", "certifications_count", "education_count"]:
        rows.append({"section": "nested_count_mean", "metric": col, "value": float(nested[col].mean())})
        rows.append({"section": "nested_count_max", "metric": col, "value": int(nested[col].max())})
        for key, value in nested[col].value_counts().sort_index().items():
            rows.append({"section": f"distribution_{col}", "metric": str(key), "value": int(value)})
    save_hist(profiles["years_of_experience"], "Years of Experience", "years_of_experience_hist.png", kde=True)
    save_box(profiles["years_of_experience"], "Years of Experience", "years_of_experience_box.png")
    for col in ["current_company_size", "current_industry", "country"]:
        save_bar(profiles[col].value_counts(), col.replace("_", " ").title(), f"{col}_bar.png")
    for col in ["career_history_count", "skills_count", "certifications_count", "education_count"]:
        save_hist(nested[col], col.replace("_", " ").title(), f"{col}_hist.png", bins=range(int(nested[col].max()) + 2))
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_DIR / "dataset_summary.csv", index=False)
    return nested


def title_analysis(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    profiles = frames["profiles"]
    careers = frames["careers"]
    current = profiles["current_title"].fillna("")
    historical = careers.loc[~careers.get("is_current", False).astype(bool), "title"].fillna("")
    current_norm = current.map(normalize_text)
    historical_norm = historical.map(normalize_text)
    word_counts = Counter()
    for title in pd.concat([current_norm, historical_norm]):
        word_counts.update(tokenize(title))
    rows = []
    for kind, series, top_n in [("current_title", current_norm, 100), ("historical_title", historical_norm, 100)]:
        counts = series.value_counts().head(top_n)
        for title, count in counts.items():
            rows.append({"section": kind, "item": title, "count": int(count), "share": count / len(series)})
    for word, count in word_counts.most_common(200):
        rows.append({"section": "title_word", "item": word, "count": int(count), "share": count / max(1, sum(word_counts.values()))})
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_DIR / "titles_summary.csv", index=False)
    save_bar(current_norm.value_counts(), "Most Common Current Titles", "current_titles_bar.png")
    save_bar(pd.Series(dict(word_counts.most_common(40))), "Most Common Title Terms", "title_terms_bar.png")
    return out


def skills_analysis(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    skills = frames["skills"].copy()
    skills["skill_norm"] = skills["name"].map(normalize_text)
    rows = []
    skill_counts = skills["skill_norm"].value_counts()
    for skill, count in skill_counts.head(500).items():
        rows.append({"section": "skill_frequency", "item": skill, "count": int(count), "share": count / skills["candidate_id"].nunique()})
    for prof, count in skills["proficiency"].value_counts().items():
        rows.append({"section": "proficiency", "item": prof, "count": int(count), "share": count / len(skills)})
    pair_counts: Counter[tuple[str, str]] = Counter()
    triplet_counts: Counter[tuple[str, str, str]] = Counter()
    for _, group in tqdm(skills.groupby("candidate_id"), desc="Counting skill combinations"):
        names = sorted(set(group["skill_norm"].dropna()))
        top_names = [n for n in names if n]
        pair_counts.update(combinations(top_names, 2))
        triplet_counts.update(combinations(top_names, 3))
    for pair, count in pair_counts.most_common(200):
        rows.append({"section": "skill_pair", "item": " | ".join(pair), "count": int(count), "share": count / skills["candidate_id"].nunique()})
    for triplet, count in triplet_counts.most_common(100):
        rows.append({"section": "skill_triplet", "item": " | ".join(triplet), "count": int(count), "share": count / skills["candidate_id"].nunique()})
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_DIR / "skills_summary.csv", index=False)
    save_bar(skill_counts, "Top Skills", "skills_bar.png", top_n=40)
    save_hist(skill_counts, "Skill Frequency Distribution", "skill_frequency_distribution.png", bins=60)
    return out


def signals_analysis(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    signals = frames["signals"].copy()
    bool_cols = signals.select_dtypes(include=["bool"]).columns
    for col in bool_cols:
        signals[col] = signals[col].astype(int)
    numeric = signals.select_dtypes(include=[np.number])
    stats = []
    for col in numeric.columns:
        if col == "candidate_id":
            continue
        s = pd.to_numeric(numeric[col], errors="coerce").dropna()
        if s.empty:
            continue
        stats.append(
            {
                "signal": col,
                "mean": s.mean(),
                "median": s.median(),
                "std": s.std(),
                "min": s.min(),
                "max": s.max(),
                "p25": s.quantile(0.25),
                "p75": s.quantile(0.75),
                "p90": s.quantile(0.90),
                "p95": s.quantile(0.95),
                "p99": s.quantile(0.99),
                "non_null_count": int(s.count()),
            }
        )
        save_hist(s, col.replace("_", " ").title(), f"signal_{col}_hist.png", kde=s.nunique() > 10)
    out = pd.DataFrame(stats)
    out.to_csv(OUTPUT_DIR / "signals_summary.csv", index=False)
    corr = numeric.drop(columns=["candidate_id"], errors="ignore").corr(numeric_only=True)
    if not corr.empty:
        plt.figure(figsize=(14, 11))
        sns.heatmap(corr, cmap="vlag", center=0)
        plt.title("Behavioral Signal Correlation Matrix")
        plt.tight_layout()
        plt.savefig(VIZ_DIR / "signals_correlation_heatmap.png", dpi=160)
        plt.close()
    relationship_report(frames, corr)
    return out


def scatter_relationship(signals: pd.DataFrame, x: str, y: str, filename: str) -> dict[str, Any]:
    clean = signals[[x, y]].apply(pd.to_numeric, errors="coerce").dropna()
    if clean.empty:
        return {"pair": f"{x} vs {y}", "n": 0, "pearson": np.nan, "spearman": np.nan}
    sample = clean.sample(n=min(10000, len(clean)), random_state=RANDOM_SEED)
    plt.figure(figsize=(7, 5))
    sns.scatterplot(data=sample, x=x, y=y, alpha=0.25, s=14, edgecolor=None)
    plt.title(f"{x} vs {y}")
    plt.tight_layout()
    plt.savefig(VIZ_DIR / filename, dpi=160)
    plt.close()
    return {
        "pair": f"{x} vs {y}",
        "n": len(clean),
        "pearson": clean[x].corr(clean[y], method="pearson"),
        "spearman": clean[x].corr(clean[y], method="spearman"),
    }


def relationship_report(frames: dict[str, pd.DataFrame], corr: pd.DataFrame) -> None:
    signals = frames["signals"].copy()
    pairs = [
        ("recruiter_response_rate", "saved_by_recruiters_30d"),
        ("github_activity_score", "search_appearance_30d"),
        ("interview_completion_rate", "offer_acceptance_rate"),
        ("profile_views_received_30d", "search_appearance_30d"),
        ("profile_views_received_30d", "saved_by_recruiters_30d"),
    ]
    rows = []
    for x, y in pairs:
        if x in signals and y in signals:
            rows.append(scatter_relationship(signals, x, y, f"relationship_{x}_vs_{y}.png"))
    strongest = []
    if corr is not None and not corr.empty:
        stacked = corr.where(np.triu(np.ones(corr.shape), k=1).astype(bool)).stack().sort_values(key=lambda s: s.abs(), ascending=False)
        strongest = [(a, b, value) for (a, b), value in stacked.head(20).items()]
    lines = ["# Behavioral Signal Relationships", ""]
    lines.append("## Requested Relationships")
    for row in rows:
        lines.append(
            f"- `{row['pair']}`: n={row['n']}, Pearson={row['pearson']:.3f}, Spearman={row['spearman']:.3f}."
        )
    lines.append("")
    lines.append("## Strongest Numeric Correlations")
    for a, b, value in strongest:
        lines.append(f"- `{a}` vs `{b}`: correlation={value:.3f}.")
    (OUTPUT_DIR / "behavioral_relationships.md").write_text("\n".join(lines), encoding="utf-8")


def behavior_segments(frames: dict[str, pd.DataFrame]) -> pd.Series:
    signals = frames["signals"].copy()
    days_since = (TODAY - signals["last_active_date"]).dt.days
    rr = signals["recruiter_response_rate"].fillna(0)
    search = signals["search_appearance_30d"].fillna(0)
    saved = signals["saved_by_recruiters_30d"].fillna(0)
    interview = signals["interview_completion_rate"].fillna(0)

    search_hi = search.quantile(0.75)
    saved_hi = saved.quantile(0.75)
    segment = pd.Series("passive", index=signals.index, dtype="object")
    segment[(days_since <= 14) & (rr >= 0.65) & ((search >= search_hi) | (saved >= saved_hi)) & (interview >= 0.70)] = "highly_active"
    segment[(days_since <= 45) & (rr >= 0.35) & (interview >= 0.45) & (segment != "highly_active")] = "moderately_active"
    segment[(days_since > 90) | ((rr < 0.15) & (search <= search.quantile(0.25)) & (saved <= saved.quantile(0.25)))] = "dormant"
    signals["behavior_segment"] = segment
    frames["signals"] = signals

    summary = signals.groupby("behavior_segment").agg(
        candidates=("candidate_id", "count"),
        median_days_since_last_active=("last_active_date", lambda s: float((TODAY - s).dt.days.median())),
        median_recruiter_response_rate=("recruiter_response_rate", "median"),
        median_search_appearance_30d=("search_appearance_30d", "median"),
        median_saved_by_recruiters_30d=("saved_by_recruiters_30d", "median"),
        median_interview_completion_rate=("interview_completion_rate", "median"),
    )
    lines = [
        "# Behavioral Signal Segmentation",
        "",
        "Exploratory buckets only. These are not scores and are not ranking labels.",
        "",
        "## Bucket Definitions",
        f"- `highly_active`: active within 14 days, recruiter response >= 0.65, high search or save volume, interview completion >= 0.70.",
        f"- `moderately_active`: active within 45 days, recruiter response >= 0.35, interview completion >= 0.45, and not highly active.",
        "- `dormant`: inactive for more than 90 days, or very low response plus bottom-quartile search/save activity.",
        "- `passive`: remaining candidates.",
        "",
        "## Segment Summary",
        summary.to_markdown(),
    ]
    (OUTPUT_DIR / "behavior_segments.md").write_text("\n".join(lines), encoding="utf-8")
    return segment


def companies_analysis(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    profiles = frames["profiles"]
    careers = frames["careers"]
    current = profiles["current_company"].map(normalize_text)
    historical = careers["company"].map(normalize_text)
    rows = []
    for section, series in [("current_company", current), ("historical_company", historical)]:
        for company, count in series.value_counts().head(200).items():
            rows.append({"section": section, "item": company, "count": int(count), "share": count / len(series)})
    for size, count in profiles["current_company_size"].value_counts().items():
        rows.append({"section": "current_company_size", "item": size, "count": int(count), "share": count / len(profiles)})
    company_text = current.fillna("")
    rows.extend(company_category_rows(company_text, "current_company_category"))
    out = pd.DataFrame(rows)
    out.to_csv(OUTPUT_DIR / "companies_summary.csv", index=False)
    save_bar(current.value_counts(), "Top Current Companies", "current_companies_bar.png", top_n=40)
    return out


def company_category_rows(companies: pd.Series, section: str) -> list[dict[str, Any]]:
    counts = Counter()
    for name in companies:
        if any(c in name for c in PRODUCT_COMPANIES):
            counts["known_product"] += 1
        elif any(c in name for c in SERVICE_COMPANIES):
            counts["known_services_or_consulting"] += 1
        elif "startup" in name:
            counts["startup_named"] += 1
        else:
            counts["uncategorized"] += 1
    total = sum(counts.values())
    return [{"section": section, "item": key, "count": value, "share": value / total} for key, value in counts.items()]


def jd_alignment(frames: dict[str, pd.DataFrame], jd_text: str) -> pd.Series:
    profiles = frames["profiles"].copy()
    careers = frames["careers"]
    skills = frames["skills"].copy()
    texts = frames["texts"].copy()
    jd_norm = normalize_text(jd_text)
    term_hits = {category: [term for term in terms if normalize_text(term) in jd_norm] for category, terms in JD_TERMS.items()}
    if not any(term_hits.values()):
        term_hits = JD_TERMS

    aligned = pd.Series(False, index=profiles["candidate_id"])
    category_counts = {}
    lines = ["# JD Alignment Discovery", "", "Exploratory only. No candidate scores or ranking labels are produced.", ""]
    lines.append("## JD Term Families")
    for category, terms in term_hits.items():
        lines.append(f"- `{category}`: {', '.join(terms)}")
    lines.append("")
    for category, terms in term_hits.items():
        terms_norm = [normalize_text(t) for t in terms]
        pattern = "|".join(re.escape(t) for t in terms_norm if t)
        if not pattern:
            continue
        text_hits = texts["search_text"].str.contains(pattern, regex=True, na=False)
        matched_ids = texts.loc[text_hits, "candidate_id"]
        aligned.loc[matched_ids] = True
        category_counts[category] = int(text_hits.sum())
        title_hits = profiles["current_title"].map(normalize_text).str.contains(pattern, regex=True, na=False)
        skill_hits = skills["skill_norm" if "skill_norm" in skills else "name"].map(normalize_text).str.contains(pattern, regex=True, na=False)
        top_titles = profiles.loc[title_hits, "current_title"].map(normalize_text).value_counts().head(15)
        top_skills = skills.loc[skill_hits, "name"].map(normalize_text).value_counts().head(20)
        lines.extend([f"## {category.replace('_', ' ').title()}", f"- Candidate text matches: {int(text_hits.sum())}"])
        if not top_titles.empty:
            lines.append("- Most common aligned current titles:")
            lines.extend([f"  - `{k}`: {v}" for k, v in top_titles.items()])
        if not top_skills.empty:
            lines.append("- Most common aligned skills:")
            lines.extend([f"  - `{k}`: {v}" for k, v in top_skills.items()])
        lines.append("")

    profiles_by_id = profiles.set_index("candidate_id")
    current_company = profiles_by_id.loc[aligned[aligned].index, "current_company"].map(normalize_text)
    industry = profiles_by_id.loc[aligned[aligned].index, "current_industry"].map(normalize_text)
    lines.append("## Career Pattern Observations")
    lines.append(f"- Profiles with at least one JD-term family match in title, summary, role descriptions, or skills: {int(aligned.sum())}.")
    lines.append("- Common industries among matched profiles:")
    lines.extend([f"  - `{k}`: {v}" for k, v in industry.value_counts().head(15).items()])
    lines.append("- Common current companies among matched profiles:")
    lines.extend([f"  - `{k}`: {v}" for k, v in current_company.value_counts().head(20).items()])
    (OUTPUT_DIR / "jd_alignment_candidates.md").write_text("\n".join(lines), encoding="utf-8")
    return aligned


def timeline_checks(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    profiles = frames["profiles"].set_index("candidate_id")
    careers = frames["careers"].copy()
    careers["duration_months"] = pd.to_numeric(careers["duration_months"], errors="coerce").fillna(0)
    careers["effective_end_date"] = careers["end_date"].fillna(TODAY)
    careers["is_current_bool"] = careers["is_current"].astype(bool)
    reasons: dict[str, list[str]] = defaultdict(list)

    total_months = careers.groupby("candidate_id")["duration_months"].sum()
    yoe_months = pd.to_numeric(profiles["years_of_experience"], errors="coerce") * 12
    too_long = total_months[total_months > yoe_months.reindex(total_months.index).fillna(np.inf) + 18]
    for cid, months in too_long.items():
        reasons[cid].append(
            f"role_duration_sum_exceeds_yoe: total_months={months:.0f}, yoe_months={yoe_months.get(cid, np.nan):.0f}"
        )

    bad_dates = careers[careers["effective_end_date"] < careers["start_date"]]
    for cid, count in bad_dates.groupby("candidate_id").size().items():
        reasons[cid].append(f"end_before_start_roles={int(count)}")

    sorted_roles = careers.sort_values(["candidate_id", "start_date", "effective_end_date"])
    sorted_roles["prev_end"] = sorted_roles.groupby("candidate_id")["effective_end_date"].shift()
    sorted_roles["gap_days"] = (sorted_roles["start_date"] - sorted_roles["prev_end"]).dt.days
    overlaps = sorted_roles[sorted_roles["gap_days"] < -31].groupby("candidate_id").size()
    for cid, count in overlaps.items():
        reasons[cid].append(f"overlap_over_31_days_roles={int(count)}")
    large_gaps = sorted_roles[sorted_roles["gap_days"] > 730].groupby("candidate_id")["gap_days"].max()
    for cid, days in large_gaps.items():
        reasons[cid].append(f"large_gap_days={int(days)}")

    current_counts = careers.groupby("candidate_id")["is_current_bool"].sum()
    for cid, count in current_counts[current_counts != 1].items():
        reasons[cid].append(f"current_role_count={int(count)}")
    current_with_end = careers[careers["is_current_bool"] & careers["end_date"].notna()].groupby("candidate_id").size()
    for cid, count in current_with_end.items():
        reasons[cid].append(f"current_role_has_end_date_count={int(count)}")

    flags = [{"candidate_id": cid, "flag_count": len(items), "reasons": "; ".join(items)} for cid, items in reasons.items()]
    out = pd.DataFrame(flags)
    if not out.empty:
        out = out.sort_values(["flag_count", "candidate_id"], ascending=[False, True])
    out.to_csv(OUTPUT_DIR / "suspicious_timeline_profiles.csv", index=False)
    return out


def skill_checks(frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    skills = frames["skills"].copy()
    profiles = frames["profiles"].set_index("candidate_id")
    skills["duration_months"] = pd.to_numeric(skills["duration_months"], errors="coerce").fillna(0)
    skills["endorsements"] = pd.to_numeric(skills["endorsements"], errors="coerce").fillna(0)
    expert_counts = skills[skills["proficiency"].eq("expert")].groupby("candidate_id").size()
    expert_threshold = max(10, int(expert_counts.quantile(0.99))) if not expert_counts.empty else 10
    flags = defaultdict(list)
    low_duration = skills[skills["proficiency"].eq("expert") & (skills["duration_months"] <= 6)]
    for cid, group in low_duration.groupby("candidate_id"):
        examples = ", ".join(f"{name}({int(months)}m)" for name, months in group[["name", "duration_months"]].head(8).itertuples(index=False))
        flags[cid].append(f"expert_low_duration:{examples}")
    zero_endorsements = skills[skills["proficiency"].eq("expert") & (skills["endorsements"] == 0)]
    for cid, group in zero_endorsements.groupby("candidate_id"):
        examples = ", ".join(group["name"].head(8).astype(str))
        flags[cid].append(f"expert_zero_endorsements:{examples}")
    for cid, count in expert_counts.items():
        if count >= expert_threshold:
            flags[cid].append(f"unusually_many_expert_skills:{int(count)}")
    skill_counts = skills.groupby("candidate_id").size()
    high_skill_threshold = int(skill_counts.quantile(0.995))
    for cid, count in skill_counts.items():
        if count >= high_skill_threshold:
            flags[cid].append(f"extreme_skill_count:{int(count)}")
    rows = []
    for cid, reasons in flags.items():
        rows.append(
            {
                "candidate_id": cid,
                "flag_count": len(reasons),
                "years_of_experience": profiles.loc[cid, "years_of_experience"] if cid in profiles.index else np.nan,
                "reasons": "; ".join(reasons[:30]),
            }
        )
    out = pd.DataFrame(rows)
    if not out.empty:
        out = out.sort_values(["flag_count", "candidate_id"], ascending=[False, True])
    out.to_csv(OUTPUT_DIR / "suspicious_skill_profiles.csv", index=False)
    return out


def honeypot_report(timeline_flags: pd.DataFrame, skill_flags: pd.DataFrame, frames: dict[str, pd.DataFrame]) -> pd.DataFrame:
    signals = frames["signals"].set_index("candidate_id")
    combined = defaultdict(list)
    for _, row in timeline_flags.iterrows():
        combined[row["candidate_id"]].append(f"timeline:{row['reasons']}")
    for _, row in skill_flags.iterrows():
        combined[row["candidate_id"]].append(f"skills:{row['reasons']}")
    if "profile_completeness_score" in signals:
        high_complete_low_activity = signals[
            (signals["profile_completeness_score"] >= 95)
            & (signals["profile_views_received_30d"] <= signals["profile_views_received_30d"].quantile(0.05))
            & (signals["last_active_date"] < TODAY - pd.Timedelta(days=120))
        ]
        for cid in high_complete_low_activity.index:
            combined[cid].append("high_profile_completeness_with_low_recent_activity")
    if {"recruiter_response_rate", "avg_response_time_hours", "interview_completion_rate"}.issubset(signals.columns):
        contradictory = signals[
            (signals["recruiter_response_rate"] >= 0.85)
            & (signals["avg_response_time_hours"] >= signals["avg_response_time_hours"].quantile(0.95))
            & (signals["interview_completion_rate"] <= 0.20)
        ]
        for cid in contradictory.index:
            combined[cid].append("high_response_rate_but_slow_response_and_low_interview_completion")
    rows = [{"candidate_id": cid, "honeypot_flag_count": len(reasons), "reasons": "; ".join(reasons[:20])} for cid, reasons in combined.items()]
    out = pd.DataFrame(rows).sort_values(["honeypot_flag_count", "candidate_id"], ascending=[False, True])
    lines = [
        "# Honeypot Investigation",
        "",
        "This report flags suspicious profiles for inspection only. It does not remove, rank, or score candidates.",
        "",
        f"- Timeline-flagged profiles: {len(timeline_flags)}",
        f"- Skill-flagged profiles: {len(skill_flags)}",
        f"- Combined honeypot candidates: {len(out)}",
        "",
        "## Most Common Honeypot Reason Fragments",
    ]
    reason_counts = Counter()
    for reasons in combined.values():
        for reason in reasons:
            reason_counts.update([reason.split(":", 1)[0]])
    lines.extend([f"- `{reason}`: {count}" for reason, count in reason_counts.most_common(20)])
    lines.append("")
    lines.append("## Top Flagged Candidate IDs")
    if not out.empty:
        lines.extend([f"- `{row.candidate_id}`: {row.honeypot_flag_count} flags" for row in out.head(30).itertuples()])
    (OUTPUT_DIR / "honeypot_investigation.md").write_text("\n".join(lines), encoding="utf-8")
    return out


def write_samples(
    candidates: list[dict[str, Any]],
    frames: dict[str, pd.DataFrame],
    suspicious_ids: set[str],
    jd_aligned: pd.Series,
) -> None:
    by_id = {c["candidate_id"]: c for c in candidates}
    profiles = frames["profiles"].copy()
    texts = frames["texts"].set_index("candidate_id")["search_text"]
    rng = np.random.default_rng(RANDOM_SEED)

    groups: dict[str, list[str]] = {}
    groups["random_candidates"] = list(rng.choice(profiles["candidate_id"], size=min(50, len(profiles)), replace=False))
    groups["high_experience_candidates"] = profiles.sort_values("years_of_experience", ascending=False)["candidate_id"].head(50).tolist()
    ai_pattern = r"\b(?:ai|ml|machine learning|deep learning|nlp|llm|data scientist|computer vision)\b"
    ai_ids = texts[texts.str.contains(ai_pattern, regex=True, na=False)].index[:50].tolist()
    groups["ai_title_candidates"] = ai_ids
    consulting_pattern = r"\b(?:consultant|consulting|deloitte|accenture|pwc|kpmg|ey|capgemini|infosys|tcs|wipro)\b"
    groups["consulting_background_candidates"] = texts[texts.str.contains(consulting_pattern, regex=True, na=False)].index[:50].tolist()
    suspicious_list = sorted(suspicious_ids)[:50]
    groups["suspicious_candidates"] = suspicious_list
    groups["jd_aligned_candidates"] = jd_aligned[jd_aligned].index[:50].tolist()

    metadata = {}
    for group_name, ids in groups.items():
        group_dir = SAMPLES_DIR / group_name
        group_dir.mkdir(exist_ok=True)
        metadata[group_name] = ids
        for cid in ids:
            if cid in by_id:
                (group_dir / f"{cid}.json").write_text(json.dumps(by_id[cid], indent=2, ensure_ascii=False), encoding="utf-8")
    (SAMPLES_DIR / "sample_metadata.json").write_text(json.dumps(metadata, indent=2), encoding="utf-8")


def archetypes(frames: dict[str, pd.DataFrame], jd_aligned: pd.Series) -> None:
    profiles = frames["profiles"].copy()
    signals = frames["signals"].copy()
    skills = frames["skills"].copy()
    texts = frames["texts"].copy()
    skill_counts = skills.assign(skill_norm=skills["name"].map(normalize_text)).pivot_table(
        index="candidate_id", columns="skill_norm", values="name", aggfunc="count", fill_value=0
    )
    top_skill_cols = skills["name"].map(normalize_text).value_counts().head(80).index
    skill_features = skill_counts.reindex(columns=top_skill_cols, fill_value=0)

    vectorizer = CountVectorizer(max_features=80, min_df=25, ngram_range=(1, 2), stop_words="english")
    title_matrix = vectorizer.fit_transform(profiles["current_title"].fillna("").map(normalize_text))
    title_df = pd.DataFrame(title_matrix.toarray(), index=profiles["candidate_id"], columns=[f"title_{t}" for t in vectorizer.get_feature_names_out()])

    numeric_cols = [
        "years_of_experience",
        "profile_completeness_score",
        "recruiter_response_rate",
        "github_activity_score",
        "search_appearance_30d",
        "saved_by_recruiters_30d",
        "interview_completion_rate",
        "profile_views_received_30d",
        "skill_assessment_count",
    ]
    base = profiles[["candidate_id", "years_of_experience"]].merge(signals, on="candidate_id", how="left")
    numeric = base.set_index("candidate_id")[[c for c in numeric_cols if c in base.columns]].apply(pd.to_numeric, errors="coerce").fillna(0)
    numeric["jd_aligned_flag"] = jd_aligned.reindex(numeric.index, fill_value=False).astype(int)
    features = pd.concat([numeric, skill_features.reindex(numeric.index, fill_value=0), title_df.reindex(numeric.index, fill_value=0)], axis=1)
    sample_index = features.sample(n=min(20000, len(features)), random_state=RANDOM_SEED).index
    X = StandardScaler(with_mean=True).fit_transform(features.loc[sample_index])

    k_values = [6, 8, 10, 12]
    best_k = 8
    best_score = -1.0
    for k in k_values:
        labels = KMeans(n_clusters=k, random_state=RANDOM_SEED, n_init=10).fit_predict(X)
        if len(set(labels)) > 1:
            score = silhouette_score(X, labels, sample_size=min(5000, len(labels)), random_state=RANDOM_SEED)
            if score > best_score:
                best_score = score
                best_k = k
    kmeans = KMeans(n_clusters=best_k, random_state=RANDOM_SEED, n_init=10)
    labels = kmeans.fit_predict(X)
    pca = PCA(n_components=2, random_state=RANDOM_SEED)
    coords = pca.fit_transform(X)
    cluster_df = pd.DataFrame({"candidate_id": sample_index, "cluster": labels, "pca_1": coords[:, 0], "pca_2": coords[:, 1]})
    cluster_df = cluster_df.merge(profiles[["candidate_id", "current_title", "current_company", "current_industry", "years_of_experience"]], on="candidate_id", how="left")
    cluster_df.to_csv(OUTPUT_DIR / "candidate_archetype_clusters.csv", index=False)

    plt.figure(figsize=(9, 7))
    sns.scatterplot(data=cluster_df, x="pca_1", y="pca_2", hue="cluster", palette="tab10", s=12, alpha=0.45)
    plt.title("Candidate Archetype PCA Projection")
    plt.tight_layout()
    plt.savefig(VIZ_DIR / "candidate_archetypes_pca.png", dpi=160)
    plt.close()

    agglom_labels = AgglomerativeClustering(n_clusters=min(best_k, 10)).fit_predict(X[: min(5000, len(X))])
    lines = [
        "# Candidate Archetypes Report",
        "",
        "Clustering is unsupervised and exploratory. Cluster IDs are not relevance tiers or candidate scores.",
        "",
        f"- Sampled profiles for clustering: {len(sample_index)}",
        f"- Selected KMeans clusters: {best_k}",
        f"- Best sampled silhouette score among tested K values: {best_score:.3f}",
        f"- Hierarchical clustering sanity sample size: {len(agglom_labels)}",
        "",
    ]
    for cluster_id, group in cluster_df.groupby("cluster"):
        lines.append(f"## Cluster {cluster_id}")
        lines.append(f"- Size: {len(group)}")
        lines.append("- Common current titles:")
        lines.extend([f"  - `{k}`: {v}" for k, v in group["current_title"].map(normalize_text).value_counts().head(8).items()])
        lines.append("- Common industries:")
        lines.extend([f"  - `{k}`: {v}" for k, v in group["current_industry"].map(normalize_text).value_counts().head(6).items()])
        lines.append("- Representative candidate IDs:")
        lines.extend([f"  - `{cid}`" for cid in group["candidate_id"].head(8)])
        lines.append("")
    (OUTPUT_DIR / "candidate_archetypes_report.md").write_text("\n".join(lines), encoding="utf-8")


def hypotheses_and_findings(
    schema_df: pd.DataFrame,
    nested: pd.DataFrame,
    frames: dict[str, pd.DataFrame],
    timeline_flags: pd.DataFrame,
    skill_flags: pd.DataFrame,
    honeypot_df: pd.DataFrame,
    jd_aligned: pd.Series,
) -> None:
    profiles = frames["profiles"]
    signals = frames["signals"]
    skills = frames["skills"]
    texts = frames["texts"]
    hypotheses = ["# Dataset Hypotheses", ""]
    findings = ["# Final Findings", ""]

    low_fill = schema_df[schema_df["fill_rate"] < 0.50].head(15)
    hypotheses.append("## Hypothesis 1: Some fields are structurally sparse and may matter more as presence indicators than raw values.")
    hypotheses.append("Evidence:")
    hypotheses.extend([f"- `{row.field}` fill_rate={row.fill_rate}" for row in low_fill.itertuples()])
    hypotheses.append("")

    ai_pattern = r"\b(?:ai|ml|machine learning|deep learning|nlp|llm|computer vision|retrieval|ranking|recommendation|embedding|search)\b"
    ai_share = texts["search_text"].str.contains(ai_pattern, regex=True, na=False).mean()
    jd_share = jd_aligned.mean()
    hypotheses.append("## Hypothesis 2: JD-adjacent terminology appears in a measurable subset of profiles.")
    hypotheses.append(f"Evidence: AI/search/retrieval/ranking term share={ai_share:.3%}; JD-family aligned share={jd_share:.3%}.")
    hypotheses.append("")

    if "behavior_segment" in signals:
        seg_counts = signals["behavior_segment"].value_counts(normalize=True)
        hypotheses.append("## Hypothesis 3: Behavioral activity separates candidates into availability-like segments.")
        hypotheses.append("Evidence:")
        hypotheses.extend([f"- `{k}`: {v:.2%}" for k, v in seg_counts.items()])
        hypotheses.append("")

    hypotheses.append("## Hypothesis 4: Honeypot-like anomalies are concentrated in timeline and skill consistency flags.")
    hypotheses.append(f"Evidence: timeline flags={len(timeline_flags)}, skill flags={len(skill_flags)}, combined honeypot flags={len(honeypot_df)}.")
    hypotheses.append("")

    findings.append(f"- Loaded candidate count: {len(profiles)}.")
    findings.append(f"- Average career-history entries: {nested['career_history_count'].mean():.2f}; average skills: {nested['skills_count'].mean():.2f}.")
    findings.append(f"- JD-family aligned exploratory profiles: {int(jd_aligned.sum())} ({jd_aligned.mean():.2%}).")
    if "behavior_segment" in signals:
        findings.append("- Behavioral segment counts:")
        findings.extend([f"  - `{k}`: {v}" for k, v in signals["behavior_segment"].value_counts().items()])
    findings.append(f"- Suspicious timeline profiles: {len(timeline_flags)}.")
    findings.append(f"- Suspicious skill profiles: {len(skill_flags)}.")
    findings.append(f"- Honeypot investigation combined flagged profiles: {len(honeypot_df)}.")
    findings.append("")
    findings.append("No candidate scores, relevance tiers, ranking model, or submission ordering were generated.")

    (OUTPUT_DIR / "dataset_hypotheses.md").write_text("\n".join(hypotheses), encoding="utf-8")
    (OUTPUT_DIR / "final_findings.md").write_text("\n".join(findings), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    if not DATASET_PATH.exists():
        raise FileNotFoundError(f"Dataset not found: {DATASET_PATH}")
    candidates = load_candidates(DATASET_PATH)
    schema_df = schema_audit(candidates)
    frames = build_frames(candidates)
    nested = dataset_overview(frames, len(candidates))
    title_analysis(frames)
    skills_analysis(frames)
    if "skill_norm" not in frames["skills"]:
        frames["skills"]["skill_norm"] = frames["skills"]["name"].map(normalize_text)
    signals_analysis(frames)
    behavior_segments(frames)
    companies_analysis(frames)
    jd_text = read_docx_text(JD_PATH)
    jd_aligned = jd_alignment(frames, jd_text)
    timeline_flags = timeline_checks(frames)
    skill_flags = skill_checks(frames)
    honeypot_df = honeypot_report(timeline_flags, skill_flags, frames)
    suspicious_ids = set(timeline_flags.get("candidate_id", pd.Series(dtype=str))) | set(skill_flags.get("candidate_id", pd.Series(dtype=str)))
    write_samples(candidates, frames, suspicious_ids, jd_aligned)
    archetypes(frames, jd_aligned)
    hypotheses_and_findings(schema_df, nested, frames, timeline_flags, skill_flags, honeypot_df, jd_aligned)
    print(f"EDA complete. Outputs written to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
